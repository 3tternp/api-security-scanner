from typing import Any, List, Dict
import logging
import tempfile
import os
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from app.api import deps
from app.db.session import get_db
from app.models.scan import ScanJob, ScanResult
from app.models.user import User
from app.schemas.scan import (
    ScanJob as ScanJobSchema,
    ScanJobSummary,
    ScanJobCreate,
    ScanResult as ScanResultSchema,
    ScanResultUpdate,
    DashboardStats,
)
from app.scanner.engine import ScannerEngine

logger = logging.getLogger(__name__)
router = APIRouter()


# ---------------------------------------------------------------------------
# DOCX report generation helper
# ---------------------------------------------------------------------------

def _severity_color(severity: str):
    """Return an RGBColor tuple for a given severity label."""
    from docx.shared import RGBColor
    s = (severity or "").lower()
    if s == "critical":
        return RGBColor(0x8B, 0x00, 0x00)   # dark red
    elif s == "high":
        return RGBColor(0xC0, 0x39, 0x2B)   # red/orange
    elif s == "medium":
        return RGBColor(0xD3, 0x9A, 0x00)   # dark yellow
    elif s == "low":
        return RGBColor(0x1A, 0x5C, 0xAD)   # blue
    else:
        return RGBColor(0x60, 0x60, 0x60)   # grey (info)


def _add_table_row(table, label: str, value: str, label_width=None):
    """Append a two-column key/value row to *table*."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    import copy

    row = table.add_row()
    cell_label = row.cells[0]
    cell_value = row.cells[1]

    cell_label.text = label
    cell_value.text = value or ""

    # Style label cell
    for para in cell_label.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)
    for para in cell_value.paragraphs:
        for run in para.runs:
            run.font.size = Pt(9)

    # Light grey background for label cell
    tc_pr = cell_label._tc.get_or_add_tcPr()
    shd = tc_pr.get_or_add_shd()
    shd.set(qn("w:fill"), "D9D9D9")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")


def generate_docx_report(scan: ScanJob, results: list, output_path: str):
    """Generate a VAPT-style DOCX report and write it to *output_path*."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX report generation") from exc

    doc = Document()

    # -----------------------------------------------------------------------
    # Page margins
    # -----------------------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -----------------------------------------------------------------------
    # COVER PAGE
    # -----------------------------------------------------------------------
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run("API Security Vulnerability Assessment Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Penetration Testing & Security Audit")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    doc.add_paragraph()

    def _kv_center(label, value):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(11)
        r_val = p.add_run(value)
        r_val.font.size = Pt(11)

    _kv_center("Target URL", scan.target_url or "N/A")
    _kv_center("Scan ID", str(scan.id))
    _kv_center("Report Date", datetime.utcnow().strftime("%B %d, %Y"))
    _kv_center("Scan Status", scan.status or "N/A")
    if scan.completed_at:
        _kv_center("Completed At", scan.completed_at.strftime("%Y-%m-%d %H:%M UTC"))

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # TABLE OF CONTENTS (static)
    # -----------------------------------------------------------------------
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Vulnerability Findings",
        "3. Methodology",
        "4. Tools Used",
        "5. Appendix",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style="List Number")
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # EXECUTIVE SUMMARY
    # -----------------------------------------------------------------------
    doc.add_heading("1. Executive Summary", level=1)

    intro = doc.add_paragraph(
        f"This report presents the findings of an automated API security vulnerability assessment "
        f"conducted against the target: {scan.target_url}. The assessment was performed using the "
        f"API Vulnerability Scanner tool and covers the OWASP API Security Top 10 categories."
    )
    intro.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Severity counts
    severity_counts: Dict[str, int] = {
        "Critical": 0, "High": 0, "Medium": 0, "Low": 0, "Info": 0
    }
    for r in results:
        sev = (r.severity or "info").capitalize()
        if sev in severity_counts:
            severity_counts[sev] += 1
        else:
            severity_counts["Info"] += 1

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Table Grid"
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = "Severity"
    hdr_cells[1].text = "Count"
    hdr_cells[2].text = "Risk Level"

    # Bold header
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    risk_map = {
        "Critical": "Immediate action required",
        "High": "Action required",
        "Medium": "Action recommended",
        "Low": "Advisory",
        "Info": "Informational",
    }

    for sev, cnt in severity_counts.items():
        row = summary_table.add_row()
        row.cells[0].text = sev
        row.cells[1].text = str(cnt)
        row.cells[2].text = risk_map.get(sev, "")
        color = _severity_color(sev)
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = color
            run.font.size = Pt(10)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(10)
        for run in row.cells[2].paragraphs[0].runs:
            run.font.size = Pt(10)

    # Total row
    total_row = summary_table.add_row()
    total_row.cells[0].text = "Total"
    total_row.cells[1].text = str(len(results))
    total_row.cells[2].text = ""
    for cell in total_row.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # VULNERABILITY FINDINGS
    # -----------------------------------------------------------------------
    doc.add_heading("2. Vulnerability Findings", level=1)

    if not results:
        doc.add_paragraph("No vulnerabilities were identified during this assessment.")
    else:
        # Group by rule_id + description to deduplicate
        seen: Dict[str, ScanResult] = {}
        unique_results = []
        for r in results:
            key = f"{r.rule_id}::{r.description}"
            if key not in seen:
                seen[key] = r
                unique_results.append(r)

        for idx, finding in enumerate(unique_results, start=1):
            sev = (finding.severity or "info").capitalize()
            color = _severity_color(sev)

            # Finding heading
            heading_para = doc.add_heading(level=2)
            heading_para.clear()
            num_run = heading_para.add_run(f"{idx}. ")
            num_run.font.size = Pt(13)
            title_run = heading_para.add_run(finding.rule_id or "Finding")
            title_run.font.size = Pt(13)
            sev_run = heading_para.add_run(f"  [{sev}]")
            sev_run.font.size = Pt(11)
            sev_run.bold = True
            sev_run.font.color.rgb = color

            # Finding detail table
            ftable = doc.add_table(rows=0, cols=2)
            ftable.style = "Table Grid"

            # Set column widths via first row (approximate)
            def _set_col_width(table, widths_cm):
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        cell.width = Cm(widths_cm[i])

            rows_data = [
                ("Vulnerability Name", finding.rule_id or ""),
                ("Severity", sev),
                ("Status", finding.status or "Open"),
                ("URL / Endpoint", f"{finding.method or ''} {finding.endpoint or ''}".strip()),
                ("Description", finding.description or ""),
                ("Impact", finding.impact or "Unknown"),
                ("CVSS Score", finding.cvss_score or "N/A"),
                ("CVSS Vector", finding.cvss_vector or "N/A"),
                ("Attack Vector", finding.attack_vector or "N/A"),
                ("Attack Complexity", finding.attack_complexity or "N/A"),
                ("Privileges Required", finding.privileges_required or "N/A"),
                ("User Interaction", finding.user_interaction or "N/A"),
                ("Scope", finding.scope or "N/A"),
                ("Confidentiality Impact", finding.confidentiality or "N/A"),
                ("Integrity Impact", finding.integrity or "N/A"),
                ("Availability Impact", finding.availability or "N/A"),
                ("Steps to Reproduce (PoC)", finding.proof_of_concept or "N/A"),
                ("Recommendation", finding.remediation or "Unknown"),
            ]

            for label, value in rows_data:
                _add_table_row(ftable, label, value)

            # Color the Severity value cell
            for row in ftable.rows:
                first_cell_text = row.cells[0].text
                if first_cell_text == "Severity":
                    val_cell = row.cells[1]
                    val_cell.paragraphs[0].clear()
                    sev_r = val_cell.paragraphs[0].add_run(sev)
                    sev_r.bold = True
                    sev_r.font.color.rgb = color
                    sev_r.font.size = Pt(9)
                    break

            doc.add_paragraph()  # spacing between findings

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # METHODOLOGY
    # -----------------------------------------------------------------------
    doc.add_heading("3. Methodology", level=1)
    method_para = doc.add_paragraph(
        "The assessment was conducted following the OWASP API Security Top 10 framework. "
        "The following categories were tested:"
    )
    method_para.runs[0].font.size = Pt(10)

    owasp_items = [
        "API1:2023 – Broken Object Level Authorization (BOLA)",
        "API2:2023 – Broken Authentication",
        "API3:2023 – Broken Object Property Level Authorization",
        "API4:2023 – Unrestricted Resource Consumption",
        "API5:2023 – Broken Function Level Authorization",
        "API6:2023 – Unrestricted Access to Sensitive Business Flows",
        "API7:2023 – Server Side Request Forgery",
        "API8:2023 – Security Misconfiguration",
        "API9:2023 – Improper Inventory Management",
        "API10:2023 – Unsafe Consumption of APIs",
    ]
    for item in owasp_items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # TOOLS USED
    # -----------------------------------------------------------------------
    doc.add_heading("4. Tools Used", level=1)
    tools = [
        ("API Vulnerability Scanner", "Custom automated scanner covering OWASP API Top 10"),
        ("python-httpx", "HTTP client used for sending test requests"),
        ("SQLAlchemy", "ORM used for scan data persistence"),
        ("FastAPI", "Web framework for scanner backend"),
        ("python-docx", "Used to generate this report"),
    ]
    tools_table = doc.add_table(rows=1, cols=2)
    tools_table.style = "Table Grid"
    tools_table.rows[0].cells[0].text = "Tool"
    tools_table.rows[0].cells[1].text = "Purpose"
    for cell in tools_table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for tool_name, purpose in tools:
        row = tools_table.add_row()
        row.cells[0].text = tool_name
        row.cells[1].text = purpose
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # APPENDIX
    # -----------------------------------------------------------------------
    doc.add_heading("5. Appendix", level=1)
    app_para = doc.add_paragraph(
        "This report was automatically generated by the API Vulnerability Scanner. "
        "All findings should be validated manually before remediation. "
        "CVSS scores are indicative and should be recalculated based on your specific environment."
    )
    app_para.runs[0].font.size = Pt(10)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# ENDPOINTS
# ---------------------------------------------------------------------------

@router.post("/", response_model=ScanJobSchema)
def create_scan(
    *,
    db: Session = Depends(get_db),
    scan_in: ScanJobCreate,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(deps.get_current_active_admin),
) -> Any:
    logger.info(f"[DEBUG] Received scan creation request: {scan_in}")
    try:
        scan = ScanJob(
            target_url=scan_in.target_url,
            spec_url=scan_in.spec_url,
            config=scan_in.config,
            status="pending",
        )
        db.add(scan)
        db.commit()
        db.refresh(scan)
        logger.info(f"[DEBUG] Scan created in DB with ID: {scan.id}")

        scanner = ScannerEngine(db, scan.id)
        background_tasks.add_task(scanner.run, scan_in.spec_content)
        logger.info(f"[DEBUG] Background task scheduled for scan {scan.id}")

        return scan
    except Exception as e:
        logger.error(f"[DEBUG] Error creating scan: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/dashboard/stats", response_model=DashboardStats)
def get_dashboard_stats(
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    """Return aggregate statistics for the dashboard."""
    try:
        all_scans = db.query(ScanJob).all()
        total_scans = len(all_scans)
        completed_scans = sum(1 for s in all_scans if s.status == "completed")
        running_scans = sum(1 for s in all_scans if s.status == "running")

        all_results = db.query(ScanResult).all()
        total_findings = len(all_results)
        open_findings = sum(1 for r in all_results if (r.status or "Open") == "Open")
        critical_findings = sum(1 for r in all_results if (r.severity or "").lower() == "critical")
        high_findings = sum(1 for r in all_results if (r.severity or "").lower() == "high")
        medium_findings = sum(1 for r in all_results if (r.severity or "").lower() == "medium")
        low_findings = sum(1 for r in all_results if (r.severity or "").lower() == "low")
        info_findings = sum(1 for r in all_results if (r.severity or "").lower() == "info")

        findings_by_rule: Dict[str, int] = {}
        for r in all_results:
            rule = r.rule_id or "unknown"
            findings_by_rule[rule] = findings_by_rule.get(rule, 0) + 1

        # Map scanner rule IDs → OWASP API Top 10 categories
        RULE_TO_OWASP: Dict[str, str] = {
            "BOLA-IDOR":        "API1",
            "AUTH-MISSING":     "API2",
            "JWT-001":          "API2",
            "SENSITIVE-DATA":   "API3",
            "MASS-ASSIGN-001":  "API3",
            "RATE-LIMIT":       "API4",
            "FUZZING":          "API4",
            "BFLA-001":         "API5",
            "BUSINESS-LOGIC":   "API6",
            "SSRF-001":         "API7",
            "SEC-HEADERS":      "API8",
            "CORS-001":         "API8",
            "HTML-INJ-001":     "API8",
            "PATH-TRAV-001":    "API8",
            "OPENAPI-CONTRACT": "API9",
            "INJECTION-BASIC":  "API10",
            "DESERIALIZATION":  "API10",
        }
        owasp_counts: Dict[str, int] = {}
        for rule_id, count in findings_by_rule.items():
            owasp_cat = RULE_TO_OWASP.get(rule_id)
            if owasp_cat:
                owasp_counts[owasp_cat] = owasp_counts.get(owasp_cat, 0) + count

        return DashboardStats(
            total_scans=total_scans,
            completed_scans=completed_scans,
            running_scans=running_scans,
            total_findings=total_findings,
            open_findings=open_findings,
            critical_findings=critical_findings,
            high_findings=high_findings,
            medium_findings=medium_findings,
            low_findings=low_findings,
            info_findings=info_findings,
            findings_by_rule=findings_by_rule,
            owasp_counts=owasp_counts,
        )
    except Exception as e:
        logger.error(f"[ERROR] Failed to get dashboard stats: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/", response_model=List[ScanJobSummary])
def read_scans(
    db: Session = Depends(get_db),
    skip: int = 0,
    limit: int = 100,
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    """Return lightweight scan summaries with finding counts (no full results payload)."""
    from sqlalchemy import func
    logger.info("DEBUG: Entering read_scans")
    try:
        # Subquery: count results per scan job
        count_subq = (
            db.query(ScanResult.job_id, func.count(ScanResult.id).label("finding_count"))
            .group_by(ScanResult.job_id)
            .subquery()
        )
        rows = (
            db.query(ScanJob, count_subq.c.finding_count)
            .outerjoin(count_subq, ScanJob.id == count_subq.c.job_id)
            .order_by(ScanJob.id.desc())
            .offset(skip)
            .limit(limit)
            .all()
        )
        summaries = []
        for scan, fc in rows:
            summaries.append(
                ScanJobSummary(
                    id=scan.id,
                    target_url=scan.target_url,
                    spec_url=scan.spec_url,
                    status=scan.status,
                    created_at=scan.created_at,
                    completed_at=scan.completed_at,
                    finding_count=fc or 0,
                )
            )
        logger.info(f"[DEBUG] Retrieved {len(summaries)} scans")
        return summaries
    except Exception as e:
        logger.error(f"[ERROR] Failed to read scans: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{scan_id}", response_model=ScanJobSchema)
def read_scan(
    scan_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    scan = db.query(ScanJob).filter(ScanJob.id == scan_id).first()
    if not scan:
        raise HTTPException(status_code=404, detail="Scan not found")
    return scan


@router.get("/{scan_id}/results", response_model=List[ScanResultSchema])
def read_scan_results(
    scan_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    logger.info(f"DEBUG: Entering read_scan_results for scan_id {scan_id}")
    try:
        results = db.query(ScanResult).filter(ScanResult.job_id == scan_id).all()
        logger.info(f"[DEBUG] Retrieved {len(results)} results")
        return results
    except Exception as e:
        logger.error(f"[ERROR] Failed to read scan results: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.patch("/results/{result_id}/status", response_model=ScanResultSchema)
def update_result_status(
    result_id: int,
    status_in: ScanResultUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    """Update the triage status of a single finding."""
    valid_statuses = {"Open", "In Progress", "Fixed", "Accepted Risk"}
    if status_in.status not in valid_statuses:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid status. Must be one of: {', '.join(sorted(valid_statuses))}",
        )
    result = db.query(ScanResult).filter(ScanResult.id == result_id).first()
    if not result:
        raise HTTPException(status_code=404, detail="Finding not found")
    try:
        result.status = status_in.status
        db.commit()
        db.refresh(result)
        return result
    except Exception as e:
        logger.error(f"[ERROR] Failed to update result status: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{scan_id}/report/docx")
def download_docx_report(
    scan_id: int,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    """Generate and return a VAPT-style DOCX report for the given scan."""
    import shutil

    scan = db.query(ScanJob).filter(ScanJob.id == scan_id).first()
    if not scan:
        raise HTTPException(status_code=404, detail="Scan not found")

    if scan.status not in ("completed", "failed"):
        raise HTTPException(
            status_code=400,
            detail="Report can only be generated for completed or failed scans.",
        )

    results = db.query(ScanResult).filter(ScanResult.job_id == scan_id).all()

    try:
        tmp_dir = tempfile.mkdtemp()
        output_path = os.path.join(tmp_dir, f"scan_{scan_id}_report.docx")
        generate_docx_report(scan, results, output_path)
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        logger.error(f"[ERROR] Failed to generate DOCX report: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Report generation failed: {e}")

    # Clean up the temp directory after the response is sent
    background_tasks.add_task(shutil.rmtree, tmp_dir, True)

    return FileResponse(
        path=output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"scan_{scan_id}_vapt_report.docx",
    )


@router.delete("/{scan_id}")
def delete_scan(
    scan_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_active_admin),
) -> Any:
    scan = db.query(ScanJob).filter(ScanJob.id == scan_id).first()
    if not scan:
        raise HTTPException(status_code=404, detail="Scan not found")

    db.query(ScanResult).filter(ScanResult.job_id == scan_id).delete()
    db.delete(scan)
    db.commit()
    return {"detail": "Scan deleted"}
